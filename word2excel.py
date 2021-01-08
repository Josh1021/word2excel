 

import win32com
from win32com.client import Dispatch, constants
import docx
from docx import Document
import xlrd
from xlrd import xldate_as_tuple
import xlwt
import os
from xlutils.copy import copy
 
def docxInfo(addr):
    document = Document(addr)
    info = {'department':[],
    'data':[],
    'title':[],
    'level':[]}

    tables = document.tables
    #  for table in document.tables:
    #     for row in table.rows:
    #         print(row)
    #         for cell in row.cells:
    #             # print(row, cell)
    #             print(cell.text)
    table_word = tables[0]
    #  print(len(table_word.rows))
    info['title'] = table_word.cell(1,0).text #标题
    for row in table_word.rows:
        for cell in row.cells:
            if "发布部门" in cell.text:
                info['department'] = cell.text.split(' ')[1]
            elif "发布日期" in cell.text:
                info['data'] = cell.text.split(' ')[1]
            elif "效力级别" in cell.text:
                info['level'] = cell.text.split(' ')[1]
    #  data = table_word.cell(2,2).text.split(' ')[1]  #发布日期
    #  department = table_word.cell(3,1).text.split(' ')[1]   #部门
    #  level= table_word.cell(3,3).text.split(' ')[1]  #级别
    print(info['title'])
    print(info['data'])
    print(info['department'])
    print(info['level'])
    return info
 
# 将 .doc 文件转成 .docx 
def doc2docx(path):
    w = win32com.client.Dispatch('Word.Application')
    w.Visible = 0
    w.DisplayAlerts = 0
    doc = w.Documents.Open(path)
    newpath = os.path.splitext(path)[0] + '.docx'
    doc.SaveAs(newpath, 12, False, "", True, "", False, False, False, False)
    try:
        print('保存文件')
        doc.Close()
        print('关闭Word.Application')
        w.Quit()
    except Exception as e:
        print(e)
    os.remove(path)
    return newpath


if __name__ == '__main__':
    # print(docxInfo(test_d))
    # memo_d = '模板.xls'
    memo_d = '模板.xls'
    memo = xlrd.open_workbook(memo_d) #读取excel
    sheet0 = memo.sheet_by_index(0) #读取�?1张表
    memo_date = sheet0.col_values(4) #读取�?5�?
    # print(memo_date)
    memo_n = len(memo_date) #去掉标题
    # # if memo_n>0:
    # #     xlsx_date = memo_date[memo_n-1] #读取最后一条记录的日期,
    # # latest_date = sheet0.cell_value(memo_n-1,5)
    # # 返回时间�?

    # # 新建一个xlsx
    memo_new = copy(memo)
    sheet1 = memo_new.get_sheet(0)

    # 重建超链�?
    # hyperlinks = sheet0.col_values(6) # xlrd读取的也是text,造成超链接丢�?
    # k = 1
    # n_hyperlink = len(hyperlinks)
    # for k in range(n_hyperlink):
    #     link = 'HYPERLINK("%s";"%s")' %(hyperlinks[k],hyperlinks[k])
    #     sheet1.write(k,6,xlwt.Formula(link))
    #     k = k+1

    # 遍历log文件夹并进行查询
    log_d ='C:\\Users\\tg\\Desktop\\xh\\土壤污染相关政策2004—2018年\\土壤污染相关政策2004—2018年\\'
    logFiles = os.listdir(log_d)
    k = 1
    for file in logFiles:
        path = log_d+  file
        if file.endswith('.doc'):
            # print(path)
            newpath = doc2docx(path)
            # print(newpath)
    for file in logFiles:
        path = log_d+ file   
        if file.endswith('.docx'):
            info = docxInfo(path) 
            sheet1.write(memo_n,0,k)
            sheet1.write(memo_n,2,info['title'])
            sheet1.write(memo_n,1,info['data'])
            sheet1.write(memo_n,3,info['department'])
            sheet1.write(memo_n,4,info['level'])
            memo_n = memo_n+1
            k += 1
    os.remove(memo_d)
    memo_new.save(memo_d)
    print('memo was updated!')